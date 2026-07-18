from __future__ import annotations

import os
from io import BytesIO
from typing import Any, Iterator

from apps.kb.kb_understanding.adapters.OcrExtractorAdapter import OcrExtractorAdapter
from apps.kb.kb_understanding.config.ExtractConfig import DEFAULT_EXTRACT_CONFIG, ExtractConfig
from apps.kb.kb_understanding.dto.ExtractPartDto import ExtractPart
from apps.kb.kb_understanding.dto.ExtractResultDto import ExtractResult
from apps.kb.kb_understanding.enums.ExtractPartType import ExtractPartType
from apps.kb.kb_understanding.enums.UnderstandingErrorCode import UnderstandingErrorCode
from apps.kb.kb_understanding.errors.UnderstandingProcessingError import UnderstandingProcessingError
from apps.kb.kb_understanding.extract.docx_metadata import (
    extract_header_footer_metadata,
    extract_merged_cells,
    extract_paragraph_metadata,
    extract_table_metadata,
)
from apps.kb.kb_understanding.extract.docx_images import iter_docx_embedded_images, open_image_blob
from apps.kb.kb_understanding.extract.extract_context import ExtractContext
from apps.kb.kb_understanding.extract.extract_limits import ExtractLimits, finalize_extract_status
from apps.kb.kb_understanding.extract.heading_path import HeadingPathTracker
from apps.kb.kb_understanding.extract.ocr_engine import OcrExtractStats
from apps.kb.kb_understanding.extract.part_builder import build_table_part, build_text_part, summarize_parts


class DocxExtractorAdapter:
    name = "python_docx"
    version = "2.3"

    def __init__(
        self,
        *,
        config: ExtractConfig | None = None,
        ocr_extractor: OcrExtractorAdapter | None = None,
    ) -> None:
        self._config = config or DEFAULT_EXTRACT_CONFIG
        self._ocr = ocr_extractor or OcrExtractorAdapter(self._config)

    def extract(self, data: bytes, *, mime_type: str | None = None) -> ExtractResult:
        return self.extract_from_bytes(data, mime_type=mime_type)

    def extract_from_path(
        self,
        path: str,
        *,
        mime_type: str | None = None,
        extract_ctx: ExtractContext | None = None,
    ) -> ExtractResult:
        limits = extract_ctx.limits if extract_ctx and extract_ctx.limits else ExtractLimits(self._config)
        limits.check_file_size_bytes(os.path.getsize(path))

        try:
            from docx import Document

            document = Document(path)
        except Exception as exc:
            raise UnderstandingProcessingError(UnderstandingErrorCode.EXTRACTION_FAILED) from exc

        return self._run_path_document(
            document,
            mime_type=mime_type,
            limits=limits,
            extract_ctx=extract_ctx,
        )

    def extract_from_bytes(
        self,
        data: bytes,
        *,
        mime_type: str | None = None,
        extract_ctx: ExtractContext | None = None,
    ) -> ExtractResult:
        from docx import Document

        limits = extract_ctx.limits if extract_ctx and extract_ctx.limits else ExtractLimits(self._config)
        limits.check_file_size(data)

        try:
            document = Document(BytesIO(data))
        except Exception as exc:
            raise UnderstandingProcessingError(UnderstandingErrorCode.EXTRACTION_FAILED) from exc

        return self._run_path_document(
            document,
            mime_type=mime_type,
            limits=limits,
            extract_ctx=extract_ctx,
        )

    def _run_path_document(
        self,
        document,
        *,
        mime_type: str | None,
        limits: ExtractLimits,
        extract_ctx: ExtractContext | None,
    ) -> ExtractResult:
        parts: list[ExtractPart] = []
        warnings: list[str] = []
        part_index = 0
        document_order = 0
        failed_blocks = 0
        table_index = 0
        heading_tracker = HeadingPathTracker()
        ocr_stats = self._resolve_ocr_stats(extract_ctx)

        def _emit(batch: list[ExtractPart]) -> None:
            if extract_ctx is not None:
                extract_ctx.emit_parts(batch, batch_size=self._config.extract_batch_size)
            else:
                parts.extend(batch)

        for section_parts in self._extract_header_footer_parts(
            document,
            start_index=part_index,
            document_order=document_order,
        ):
            if extract_ctx is not None:
                limits.check_part_count(extract_ctx.counters.total_parts + len(section_parts))
            _emit(section_parts)
            if section_parts:
                part_index = max(part.part_index for part in section_parts) + 1
                document_order = max(part.metadata.get("document_order", document_order) for part in section_parts) + 1

        block_index = 0
        for block_kind, block in self._iter_document_blocks(document):
            try:
                limits.check_duration()
            except UnderstandingProcessingError as exc:
                if exc.code == UnderstandingErrorCode.EXTRACTION_TIMEOUT.value:
                    warnings.append("extract_timeout")
                    break
                raise

            block_index += 1
            try:
                if block_kind == "paragraph":
                    block_parts = self._extract_paragraph_block(
                        block,
                        block_index=block_index,
                        start_index=part_index,
                        document_order=document_order,
                        limits=limits,
                        heading_tracker=heading_tracker,
                    )
                else:
                    block_parts = self._extract_table_block(
                        block,
                        block_index=block_index,
                        start_index=part_index,
                        document_order=document_order,
                        table_index=table_index,
                        heading_tracker=heading_tracker,
                    )
                    if block_parts:
                        table_index += 1
            except UnderstandingProcessingError:
                raise
            except Exception as exc:
                failed_blocks += 1
                warnings.append("docx_part_parse_error")
                block_parts = [
                    ExtractPart(
                        part_type=ExtractPartType.UNKNOWN.value,
                        page_number=block_index,
                        part_index=part_index,
                        text=None,
                        char_count=0,
                        status="failed",
                        error_code=UnderstandingErrorCode.DOCX_PART_PARSE_ERROR.value,
                        error_message=str(exc)[:1000],
                        metadata={
                            "source": f"docx_{block_kind}",
                            "document_order": document_order,
                            "block_kind": "unknown",
                        },
                    )
                ]

            if extract_ctx is not None:
                limits.check_part_count(extract_ctx.counters.total_parts + len(block_parts))

            _emit(block_parts)
            if block_parts:
                part_index = max(part.part_index for part in block_parts) + 1
                document_order = max(part.metadata.get("document_order", document_order) for part in block_parts) + 1

        image_parts = self._extract_embedded_image_ocr_parts(
            document,
            start_index=part_index,
            document_order=document_order,
            ocr_stats=ocr_stats,
        )
        if image_parts:
            if extract_ctx is not None:
                limits.check_part_count(extract_ctx.counters.total_parts + len(image_parts))
            _emit(image_parts)

        return self._finalize_result(
            parts=parts,
            warnings=warnings,
            section_index=block_index,
            failed_blocks=failed_blocks,
            mime_type=mime_type,
            extract_ctx=extract_ctx,
            ocr_stats=ocr_stats,
        )

    def _extract_header_footer_parts(
        self,
        document,
        *,
        start_index: int,
        document_order: int,
    ) -> Iterator[list[ExtractPart]]:
        part_index = start_index
        order = document_order
        for section_index, section in enumerate(document.sections, start=1):
            for block_kind, container, part_type in (
                ("header", section.header, ExtractPartType.HEADER),
                ("footer", section.footer, ExtractPartType.FOOTER),
            ):
                for paragraph in container.paragraphs:
                    text = (paragraph.text or "").strip()
                    if not text:
                        continue
                    metadata = extract_header_footer_metadata(
                        source=f"docx_{block_kind}",
                        text=text,
                        document_order=order,
                        part_index=part_index,
                        section_index=section_index,
                        block_kind=block_kind,
                    )
                    yield [
                        ExtractPart(
                            part_type=part_type.value,
                            page_number=section_index,
                            part_index=part_index,
                            text=text,
                            char_count=len(text),
                            metadata=metadata,
                        )
                    ]
                    part_index += 1
                    order += 1

    @staticmethod
    def _iter_document_blocks(document) -> Iterator[tuple[str, Any]]:
        from docx.document import Document as DocumentClass
        from docx.oxml.table import CT_Tbl
        from docx.oxml.text.paragraph import CT_P
        from docx.table import Table
        from docx.text.paragraph import Paragraph

        if not isinstance(document, DocumentClass):
            raise TypeError("Expected python-docx Document")

        for child in document.element.body.iterchildren():
            if isinstance(child, CT_P):
                yield "paragraph", Paragraph(child, document)
            elif isinstance(child, CT_Tbl):
                yield "table", Table(child, document)

    def _extract_paragraph_block(
        self,
        paragraph,
        *,
        block_index: int,
        start_index: int,
        document_order: int,
        limits: ExtractLimits,
        heading_tracker: HeadingPathTracker,
    ) -> list[ExtractPart]:
        text = (paragraph.text or "").strip()
        if not text:
            return []

        base_metadata = extract_paragraph_metadata(
            paragraph,
            document_order=document_order,
            part_index=start_index,
            block_index=block_index,
        )
        if base_metadata.get("is_heading") and base_metadata.get("heading_level") is not None:
            path_info = heading_tracker.update(int(base_metadata["heading_level"]), text)
        else:
            path_info = heading_tracker.current()

        parts: list[ExtractPart] = []
        index = start_index
        order = document_order
        for chunk in self._split_text_chunks(text, limits):
            metadata = dict(base_metadata)
            metadata.update(path_info)
            metadata["part_index"] = index
            metadata["document_order"] = order
            parts.append(
                build_text_part(
                    page_number=block_index,
                    part_index=index,
                    text=chunk,
                    metadata=metadata,
                )
            )
            index += 1
            order += 1
        return parts

    def _extract_table_block(
        self,
        table,
        *,
        block_index: int,
        start_index: int,
        document_order: int,
        table_index: int,
        heading_tracker: HeadingPathTracker,
    ) -> list[ExtractPart]:
        rows: list[list[str]] = []
        for row in table.rows:
            cells = [(cell.text or "").strip() for cell in row.cells]
            if any(cells):
                rows.append(cells)
        if not rows:
            return []

        headers = rows[0]
        body_rows = rows[1:] if len(rows) > 1 else []
        merged_cells = extract_merged_cells(table)
        metadata = extract_table_metadata(
            table,
            document_order=document_order,
            part_index=start_index,
            block_index=block_index,
            table_index=table_index,
            headers=headers,
            rows=body_rows,
            merged_cells=merged_cells,
        )
        metadata.update(heading_tracker.current())
        part = build_table_part(
            page_number=block_index,
            part_index=start_index,
            headers=headers,
            rows=body_rows,
            source="docx_table",
            metadata=metadata,
        )
        return [part]

    def _split_text_chunks(self, text: str, limits: ExtractLimits) -> list[str]:
        max_size = self._config.max_part_size
        if len(text) <= max_size:
            return [text]

        chunks: list[str] = []
        remaining = text
        while remaining:
            if len(remaining) <= max_size:
                chunks.append(remaining)
                break
            segment = remaining[:max_size]
            split_at = segment.rfind("\n")
            if split_at <= 0:
                split_at = segment.rfind(" ")
            if split_at <= 0:
                split_at = max_size
            piece = remaining[:split_at].strip()
            if piece:
                chunks.append(piece)
            remaining = remaining[split_at:].lstrip()
            limits.check_duration()
        return chunks

    def _resolve_ocr_stats(self, extract_ctx: ExtractContext | None) -> OcrExtractStats:
        stats = extract_ctx.ocr_stats if extract_ctx is not None else OcrExtractStats()
        stats.ocr_engine_available = self._ocr.is_available
        stats.ocr_language = self._config.ocr_languages if self._config.ocr_enabled else ""
        return stats

    @staticmethod
    def _ocr_result_fields(ocr_stats: OcrExtractStats) -> dict[str, int | bool | str]:
        return {
            "pdf_pages_ocr_scanned": ocr_stats.pdf_pages_ocr_scanned,
            "docx_images_ocr_scanned": ocr_stats.docx_images_ocr_scanned,
            "ocr_engine_available": ocr_stats.ocr_engine_available,
            "ocr_language": ocr_stats.ocr_language,
        }

    def _extract_embedded_image_ocr_parts(
        self,
        document,
        *,
        start_index: int,
        document_order: int,
        ocr_stats: OcrExtractStats,
    ) -> list[ExtractPart]:
        if not self._config.ocr_enabled or not self._config.ocr_run_on_docx_images:
            return []

        parts: list[ExtractPart] = []
        index = start_index
        order = document_order

        for image_name, blob in iter_docx_embedded_images(document):
            ocr_stats.docx_images_ocr_scanned += 1
            try:
                image = open_image_blob(blob)
                ocr_part = self._ocr.ocr_embedded_image(
                    image,
                    part_index=index,
                    document_order=order,
                    image_name=image_name,
                )
            except Exception as exc:
                ocr_part = ExtractPart(
                    part_type=ExtractPartType.OCR_FAILED.value,
                    page_number=None,
                    part_index=index,
                    text=None,
                    char_count=0,
                    status="failed",
                    error_code=UnderstandingErrorCode.OCR_FAILED.value,
                    error_message=str(exc)[:1000],
                    metadata={
                        "source": "docx_embedded_image_ocr",
                        "document_order": order,
                        "block_kind": "ocr_text",
                        "image_name": image_name,
                        "order_confidence": "low",
                    },
                )

            if ocr_part is None:
                continue

            if ocr_part.metadata.get("document_order") is None:
                ocr_part.metadata["document_order"] = order
            if "order_confidence" not in ocr_part.metadata:
                ocr_part.metadata["order_confidence"] = "low"

            parts.append(ocr_part)
            index += 1
            order += 1

        return parts

    def _finalize_result(
        self,
        *,
        parts: list[ExtractPart],
        warnings: list[str],
        section_index: int,
        failed_blocks: int,
        mime_type: str | None,
        extract_ctx: ExtractContext | None,
        ocr_stats: OcrExtractStats | None = None,
    ) -> ExtractResult:
        if extract_ctx is not None:
            extract_ctx.flush()

        source_mime = mime_type or "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        stats = ocr_stats or self._resolve_ocr_stats(extract_ctx)
        ocr_fields = self._ocr_result_fields(stats)
        status = finalize_extract_status(
            parts=parts,
            failed_pages=failed_blocks,
            warnings=warnings,
            counters=extract_ctx.counters if extract_ctx is not None else None,
        )

        if extract_ctx is not None and extract_ctx.streaming:
            return ExtractResult.from_counters(
                counters=extract_ctx.counters,
                total_pages=max(section_index, 1),
                processed_pages=max(section_index, 1),
                failed_pages=failed_blocks,
                warnings=warnings,
                status=status,
                extractor_name=self.name,
                extractor_version=self.version,
                source_mime=source_mime,
                **ocr_fields,
            )

        return ExtractResult(
            total_pages=max(section_index, 1),
            parts=parts,
            total_chars=summarize_parts(parts),
            warnings=warnings,
            status=status,
            extractor_name=self.name,
            extractor_version=self.version,
            processed_pages=max(section_index, 1),
            failed_pages=failed_blocks,
            source_mime=source_mime,
            text_parts_count=sum(
                1
                for part in parts
                if part.part_type in {ExtractPartType.TEXT.value, ExtractPartType.HEADER.value, ExtractPartType.FOOTER.value}
            ),
            table_parts_count=sum(1 for part in parts if part.part_type == ExtractPartType.TABLE.value),
            ocr_text_parts_count=sum(1 for part in parts if part.part_type == ExtractPartType.OCR_TEXT.value),
            ocr_empty_parts_count=sum(1 for part in parts if part.part_type == ExtractPartType.OCR_EMPTY.value),
            ocr_failed_parts_count=sum(1 for part in parts if part.part_type == ExtractPartType.OCR_FAILED.value),
            **ocr_fields,
        )


__all__ = ["DocxExtractorAdapter"]
