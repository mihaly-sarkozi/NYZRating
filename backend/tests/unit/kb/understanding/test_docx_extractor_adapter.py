from __future__ import annotations

import os
import tempfile
from io import BytesIO
from unittest.mock import patch

import pytest
from docx import Document

from apps.kb.kb_understanding.adapters.DocxExtractorAdapter import DocxExtractorAdapter
from apps.kb.kb_understanding.config.ExtractConfig import ExtractConfig
from apps.kb.kb_understanding.enums.ExtractPartType import ExtractPartType
from apps.kb.kb_understanding.enums.ExtractStatus import ExtractStatus
from apps.kb.kb_understanding.enums.UnderstandingErrorCode import UnderstandingErrorCode


def _save_docx(build) -> str:
    document = Document()
    build(document)
    handle = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    handle.close()
    document.save(handle.name)
    return handle.name


@pytest.fixture
def adapter() -> DocxExtractorAdapter:
    return DocxExtractorAdapter(config=ExtractConfig(max_part_size=100))


def test_extract_from_path_preserves_block_order(adapter: DocxExtractorAdapter) -> None:
    def build(document: Document) -> None:
        document.add_paragraph("Első bekezdés")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Oszlop"
        table.cell(0, 1).text = "Érték"
        table.cell(1, 0).text = "A"
        table.cell(1, 1).text = "1"
        document.add_paragraph("Második bekezdés")

    path = _save_docx(build)
    try:
        result = adapter.extract_from_path(path)
    finally:
        os.unlink(path)

    assert result.status == ExtractStatus.COMPLETED.value
    assert [part.part_type for part in result.parts] == [
        ExtractPartType.TEXT.value,
        ExtractPartType.TABLE.value,
        ExtractPartType.TEXT.value,
    ]
    assert result.parts[0].text == "Első bekezdés"
    assert result.parts[1].raw_payload["source"] == "docx_table"
    assert result.parts[1].raw_payload["headers"] == ["Oszlop", "Érték"]
    assert result.parts[2].text == "Második bekezdés"


def test_extract_from_path_splits_long_paragraph(adapter: DocxExtractorAdapter) -> None:
    long_text = "A" * 250

    def build(document: Document) -> None:
        document.add_paragraph(long_text)

    path = _save_docx(build)
    try:
        result = adapter.extract_from_path(path)
    finally:
        os.unlink(path)

    text_parts = [part for part in result.parts if part.part_type == ExtractPartType.TEXT.value]
    assert len(text_parts) >= 3
    assert sum(len(part.text or "") for part in text_parts) == 250
    assert all(len(part.text or "") <= adapter._config.max_part_size for part in text_parts)


def test_extract_from_path_does_not_delegate_to_extract_from_bytes(adapter: DocxExtractorAdapter) -> None:
    path = _save_docx(lambda document: document.add_paragraph("Csak path"))
    try:
        with patch.object(adapter, "extract_from_bytes") as bytes_mock:
            result = adapter.extract_from_path(path)
    finally:
        os.unlink(path)

    bytes_mock.assert_not_called()
    assert result.parts[0].text == "Csak path"


def test_extract_from_path_partial_on_table_parse_error(adapter: DocxExtractorAdapter) -> None:
    def build(document: Document) -> None:
        document.add_paragraph("Marad")
        document.add_table(rows=1, cols=1).cell(0, 0).text = "X"

    path = _save_docx(build)
    with patch.object(
        DocxExtractorAdapter,
        "_extract_table_block",
        side_effect=RuntimeError("broken table"),
    ):
        try:
            result = adapter.extract_from_path(path)
        finally:
            os.unlink(path)

    assert result.status == ExtractStatus.PARTIAL.value
    assert any(part.part_type == ExtractPartType.TEXT.value for part in result.parts)
    failed = [part for part in result.parts if part.status == "failed"]
    assert failed
    assert failed[0].error_code == UnderstandingErrorCode.DOCX_PART_PARSE_ERROR.value


def test_extract_from_bytes_still_works_for_small_files() -> None:
    document = Document()
    document.add_paragraph("Bytes mód")
    buffer = BytesIO()
    document.save(buffer)

    adapter = DocxExtractorAdapter()
    result = adapter.extract_from_bytes(buffer.getvalue())

    assert result.status == ExtractStatus.COMPLETED.value
    assert any(part.text == "Bytes mód" for part in result.parts)


def test_extract_from_path_sets_paragraph_style_metadata(adapter: DocxExtractorAdapter) -> None:
    def build(document: Document) -> None:
        title = document.add_paragraph("Cím")
        title.style = "Heading 1"

    path = _save_docx(build)
    try:
        result = adapter.extract_from_path(path)
    finally:
        os.unlink(path)

    assert result.parts[0].metadata["source"] == "docx_paragraph"
    assert result.parts[0].metadata["style_name"] == "Heading 1"
    assert result.parts[0].metadata["block_kind"] == "heading"
    assert result.parts[0].metadata["heading_level"] == 1
