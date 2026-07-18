from __future__ import annotations

from io import BytesIO

import pytest
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase.pdfmetrics import stringWidth
from reportlab.pdfgen import canvas

from shared.documents import extract_document_from_upload
from shared.documents.pdf_layout_parser import (
    _detect_repeated_edge_lines,
    _looks_like_lonely_clause_number_line,
    _looks_like_noise_line,
    _PdfLine,
)

pytestmark = [pytest.mark.unit, pytest.mark.must_pass]


def _build_sample_pdf() -> bytes:
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4

    pdf.setFont("Helvetica-Bold", 16)
    pdf.drawString(72, height - 72, "Pilot document")

    pdf.setFont("Helvetica", 11)
    body_line = "This is the first paragraph and it should stay together as body content."
    pdf.drawString(72, height - 110, body_line)
    pdf.drawString(72, height - 126, "The second line belongs to the same paragraph and continues the thought.")

    pdf.drawString(90, height - 165, "1. First checklist item in the list")
    pdf.drawString(90, height - 181, "2. Second checklist item in the list")
    pdf.drawString(72, height - 220, "Name")
    pdf.drawString(250, height - 220, "Qty")
    pdf.drawString(400, height - 220, "Price")

    footer = "page 1"
    pdf.setFont("Helvetica", 9)
    pdf.drawString(width - 72 - stringWidth(footer, "Helvetica", 9), 28, footer)

    pdf.showPage()
    pdf.save()
    return buffer.getvalue()


def _build_hierarchical_list_pdf() -> bytes:
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    _, height = A4

    pdf.setFont("Helvetica", 11)
    pdf.drawString(90, height - 90, "12.2. The principal authorizes the broker to record the data")
    pdf.drawString(114, height - 106, "for automated newsletter service in compliance with the applicable act,")
    pdf.drawString(114, height - 122, "and to send commercial and financial updates to the client.")
    pdf.drawString(90, height - 154, "12.3. The principal may revoke this authorization in writing.")

    pdf.showPage()
    pdf.save()
    return buffer.getvalue()


def _build_docx_with_header() -> bytes:
    document = Document()
    title = document.add_paragraph()
    title.style = "Heading 1"
    title.add_run("Kockázatkezelési feltételek")
    body = document.add_paragraph()
    body.add_run("A biztosító a következő feltételek szerint jár el.")
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _build_table_of_contents_pdf() -> bytes:
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    _, height = A4

    pdf.setFont("Helvetica-Bold", 15)
    pdf.drawString(72, height - 72, "Tartalomjegyzék")
    pdf.setFont("Helvetica", 11)
    pdf.drawString(72, height - 110, "1. Bevezetés ........ 3")
    pdf.drawString(72, height - 126, "2. Díjtételek ........ 5")

    pdf.showPage()
    pdf.save()
    return buffer.getvalue()


def _build_docx_with_table() -> bytes:
    document = Document()
    document.add_paragraph("Tartalomjegyzék")
    document.add_paragraph("1. Bevezetés ........ 2")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Mező"
    table.cell(0, 1).text = "Érték"
    table.cell(1, 0).text = "Bonus-malus"
    table.cell(1, 1).text = "A10"
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _build_pdf_with_indented_paragraph() -> bytes:
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    _, height = A4

    pdf.setFont("Helvetica", 11)
    pdf.drawString(72, height - 90, "This paragraph starts normally and should remain one block even if the next line is indented.")
    pdf.drawString(108, height - 108, "The continuation line is intentionally indented, but it still belongs to the same paragraph-like block.")

    pdf.showPage()
    pdf.save()
    return buffer.getvalue()


def _build_pdf_with_blank_line_between_paragraphs() -> bytes:
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    _, height = A4

    pdf.setFont("Helvetica", 11)
    pdf.drawString(72, height - 90, "This is the first paragraph on the page.")
    pdf.drawString(72, height - 106, "It has a normal wrapped continuation line.")
    pdf.drawString(72, height - 146, "This should start a new paragraph after one blank line.")
    pdf.drawString(72, height - 162, "It also has a continuation line.")

    pdf.showPage()
    pdf.save()
    return buffer.getvalue()


def _build_pdf_with_wrapped_legal_clause() -> bytes:
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    _, height = A4

    pdf.setFont("Helvetica-Bold", 13)
    pdf.drawString(72, height - 72, "7.E. A biztosítási díj módosítása")

    pdf.setFont("Helvetica", 11)
    pdf.drawString(90, height - 102, "7.E.1. A biztosítási szerződés hatálya alatt a biztosító a biztosítás díjának mértékét az alábbiakban meghatáro-")
    pdf.drawString(102, height - 118, "zott esetekben módosíthatja.")

    pdf.showPage()
    pdf.save()
    return buffer.getvalue()


def _build_pdf_with_lowercase_numbered_items() -> bytes:
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    _, height = A4

    pdf.setFont("Helvetica", 11)
    pdf.drawString(
        72,
        height - 90,
        "5. a buntetoeljaras soran a nyomozo hatosag, ugyesz vagy a birosag altal hozott hatarozatot,",
    )
    pdf.drawString(
        90,
        height - 106,
        "tovabba az eljaras soran keletkezett szakertoi velemenyeket, jegyzokonyveket es tanunyilatkozatokat,",
    )
    pdf.drawString(72, height - 128, "6. a biztositottnak a biztositasi esemenyhez kapcsolodo dokumentumait,")
    pdf.drawString(90, height - 144, "kulonosen a korabbi orvosi iratokat es igazolasokat.")

    pdf.save()
    return buffer.getvalue()


def _build_pdf_with_roman_and_dash_markers() -> bytes:
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    _, height = A4

    pdf.setFont("Helvetica", 11)
    pdf.drawString(72, height - 90, "Elso bekezdes bevezeto szovege ami meg egy blokk.")
    pdf.drawString(72, height - 108, "IV. romai szammal kezdodo kulon pont.")
    pdf.drawString(72, height - 126, "- gondolatjellel kezdodo kulon pont.")

    pdf.save()
    return buffer.getvalue()


def _build_pdf_with_wide_gap_sentence() -> bytes:
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    _, height = A4

    pdf.setFont("Helvetica", 11)
    pdf.drawString(72, height - 90, "A biztosito a szerzodes fennallasa alatt is")
    pdf.drawString(360, height - 90, "kotelezetseget teljesiti a kartalanitas soran.")

    pdf.showPage()
    pdf.save()
    return buffer.getvalue()


def _build_pdf_cross_page_incomplete_then_place_date_stamp() -> bytes:
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    _, height = A4

    pdf.setFont("Helvetica", 11)
    pdf.drawString(72, height - 90, "A felek megallapodnak abban, hogy a szerzodes feltetelei")
    pdf.showPage()
    pdf.setFont("Helvetica", 11)
    pdf.drawString(72, height - 90, "Budapest, 2024. 03. 15.")
    pdf.save()
    return buffer.getvalue()


def _build_pdf_cross_page_title_case_continuation() -> bytes:
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    _, height = A4

    pdf.setFont("Helvetica", 11)
    # No sentence-final punctuation: layout continues on the next page.
    pdf.drawString(72, height - 90, "The insurer shall keep records of the claim and shall state that")
    pdf.showPage()
    pdf.setFont("Helvetica", 11)
    # Title Case — without cross-page merge this is often misclassified as a heading at page top.
    pdf.drawString(72, height - 90, "Further Processing Steps Still Apply In This Matter Today.")
    pdf.save()
    return buffer.getvalue()


def _build_docx_with_list_items() -> bytes:
    document = Document()
    document.add_paragraph("Bevezetes")
    document.add_paragraph("- Elso listaelem a szabalyzatban")
    document.add_paragraph("2. Masodik listaelem kulon sorban")
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def test_extract_document_from_pdf_keeps_layout_blocks() -> None:
    raw = _build_sample_pdf()

    result = extract_document_from_upload("sample.pdf", raw)

    assert result.metadata["source_format"] == "pdf"
    assert result.metadata["extraction_engine"] == "pdfplumber_layout_v1"
    assert result.paragraphs
    assert result.paragraphs[0].block_type == "heading"
    assert "Pilot document" in result.paragraphs[0].text
    assert "The second line belongs to the same paragraph" in result.paragraphs[0].text
    assert any(paragraph.block_type == "list_item" for paragraph in result.paragraphs)
    assert any(paragraph.block_type == "table_row" for paragraph in result.paragraphs)
    assert "page 1" not in result.text_content.lower()
    list_items = [paragraph.text for paragraph in result.paragraphs if paragraph.block_type == "list_item"]
    assert list_items == [
        "1. First checklist item in the list",
        "2. Second checklist item in the list",
    ]


def test_extract_document_from_pdf_returns_page_metadata_for_paragraphs() -> None:
    raw = _build_sample_pdf()

    result = extract_document_from_upload("sample.pdf", raw)

    paragraph = result.paragraphs[0]
    assert paragraph.page_number == 1
    assert paragraph.bbox is not None
    assert paragraph.font_size is not None


def test_extract_document_from_pdf_keeps_hierarchical_numbered_item_together() -> None:
    raw = _build_hierarchical_list_pdf()

    result = extract_document_from_upload("legal.pdf", raw)

    list_items = [paragraph.text for paragraph in result.paragraphs if paragraph.block_type == "list_item"]
    assert list_items == [
        "12.2. The principal authorizes the broker to record the data for automated newsletter service in compliance with the applicable act, and to send commercial and financial updates to the client.",
        "12.3. The principal may revoke this authorization in writing.",
    ]


def test_extract_document_from_pdf_keeps_indented_paragraph_together() -> None:
    raw = _build_pdf_with_indented_paragraph()

    result = extract_document_from_upload("indented.pdf", raw)

    paragraph_blocks = [paragraph for paragraph in result.paragraphs if paragraph.block_type == "paragraph"]
    assert len(paragraph_blocks) == 1
    assert "This paragraph starts normally" in paragraph_blocks[0].text
    assert "The continuation line is intentionally indented" in paragraph_blocks[0].text


def test_extract_document_from_pdf_splits_on_blank_line() -> None:
    raw = _build_pdf_with_blank_line_between_paragraphs()

    result = extract_document_from_upload("blank-line.pdf", raw)

    paragraph_blocks = [paragraph for paragraph in result.paragraphs if paragraph.block_type == "paragraph"]
    assert len(paragraph_blocks) == 2
    assert "first paragraph on the page" in paragraph_blocks[0].text
    assert "new paragraph after one blank line" in paragraph_blocks[1].text


def test_extract_document_from_pdf_keeps_wrapped_legal_list_item_in_same_block() -> None:
    raw = _build_pdf_with_wrapped_legal_clause()

    result = extract_document_from_upload("legal-wrap.pdf", raw)

    assert len(result.paragraphs) == 1
    assert result.paragraphs[0].block_type == "heading"
    assert "7.E. A biztosítási díj módosítása" in result.paragraphs[0].text
    assert "7.E.1." in result.paragraphs[0].text
    assert "meghatáro-" in result.paragraphs[0].text
    assert "módosíthatja." in result.paragraphs[0].text


def test_extract_document_from_pdf_splits_lowercase_numbered_items() -> None:
    raw = _build_pdf_with_lowercase_numbered_items()

    result = extract_document_from_upload("lowercase-numbered-items.pdf", raw)

    list_items = [paragraph.text for paragraph in result.paragraphs if paragraph.block_type == "list_item"]
    assert len(list_items) == 2
    assert list_items[0].startswith("5. a buntetoeljaras soran")
    assert "tanunyilatkozatokat" in list_items[0]
    assert list_items[1].startswith("6. a biztositottnak")
    assert "orvosi iratokat" in list_items[1]


def test_extract_document_from_pdf_splits_on_roman_and_dash_markers() -> None:
    raw = _build_pdf_with_roman_and_dash_markers()

    result = extract_document_from_upload("roman-dash-markers.pdf", raw)

    assert len(result.paragraphs) == 3
    assert result.paragraphs[0].text.startswith("Elso bekezdes")
    assert result.paragraphs[1].text.startswith("IV.")
    assert result.paragraphs[2].text.startswith("-")


def test_extract_document_from_pdf_keeps_sentence_with_large_gap_as_paragraph() -> None:
    raw = _build_pdf_with_wide_gap_sentence()

    result = extract_document_from_upload("wide-gap.pdf", raw)

    assert len(result.paragraphs) == 1
    assert result.paragraphs[0].block_type == "paragraph"
    assert result.paragraphs[0].metadata["table_cells"] == []
    assert "kotelezetseget teljesiti" in result.paragraphs[0].text


def test_extract_document_from_pdf_keeps_place_date_stamp_in_paragraph_when_only_core_blocks_split() -> None:
    raw = _build_pdf_cross_page_incomplete_then_place_date_stamp()

    result = extract_document_from_upload("place-date-stamp.pdf", raw)

    paras = [paragraph for paragraph in result.paragraphs if paragraph.block_type == "paragraph"]
    assert len(paras) == 1
    assert "szerzodes feltetelei" in paras[0].text
    assert "Budapest" in paras[0].text
    assert paras[0].metadata.get("page_span") == [1, 2]


def _build_pdf_en_dash_then_title_case_next_line() -> bytes:
    """Gondolatjel (en dash) a sor végén: a következő sor nem önálló cím, nagyobb sorköz mellett sem."""
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    _, height = A4

    pdf.setFont("Helvetica", 11)
    pdf.drawString(72, height - 90, "The insurer shall verify the claim and note that\u2013")
    pdf.drawString(72, height - 112, "Further Processing Steps Apply In This Matter Today")
    pdf.save()
    return buffer.getvalue()


def _build_pdf_ellipsis_then_title_case_next_line() -> bytes:
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    _, height = A4

    pdf.setFont("Helvetica", 11)
    pdf.drawString(72, height - 90, "The insurer shall verify the claim and document all further steps...")
    pdf.drawString(72, height - 118, "Additional Title Case Words Follow Here Today")
    pdf.save()
    return buffer.getvalue()


def _build_pdf_same_page_incomplete_then_title_case_false_heading() -> bytes:
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    _, height = A4

    pdf.setFont("Helvetica", 11)
    pdf.drawString(72, height - 90, "A biztosito a kart rendezesekor figyelembe veszi hogy")
    pdf.drawString(72, height - 114, "Tovabbi Feldolgozasi Lepesek Vonatkoznak Erre Az Ugyre")
    pdf.save()
    return buffer.getvalue()


def _build_pdf_real_legal_heading_followed_by_wrapped_paragraph() -> bytes:
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    _, height = A4

    pdf.setFont("Helvetica-Bold", 13)
    pdf.drawString(72, height - 72, "9.C. A kockazatviseles idotartama")

    pdf.setFont("Helvetica", 11)
    pdf.drawString(
        72,
        height - 102,
        "Az idojaras-biztositas kulonos szerzodesi felteteleiben meghatarozott naptari idoszakok keretein belul, a bizto-",
    )
    pdf.drawString(
        72,
        height - 118,
        "sitas aktivalasa soran a szerzodo altal az udules, kirandulas idoszokakent megjelolt egybefuggo idotartam. A",
    )
    pdf.drawString(
        72,
        height - 134,
        "kockazatviselesi idoszak kezdete nem lehet korabbi, mint a biztositas aktivalasanak a napjatol szamitott 16. nap",
    )
    pdf.drawString(72, height - 150, "07:00 oraja.")
    pdf.save()
    return buffer.getvalue()


def test_extract_document_from_pdf_keeps_en_dash_continuation_with_false_heading_next_line() -> None:
    raw = _build_pdf_en_dash_then_title_case_next_line()

    result = extract_document_from_upload("en-dash-flow.pdf", raw)

    paras = [p for p in result.paragraphs if p.block_type == "paragraph"]
    assert len(paras) == 1
    assert "\u2013" in paras[0].text or "note that" in paras[0].text
    assert "Further Processing Steps" in paras[0].text


def test_extract_document_from_pdf_keeps_same_page_false_heading_as_paragraph_continuation() -> None:
    raw = _build_pdf_same_page_incomplete_then_title_case_false_heading()

    result = extract_document_from_upload("same-page-false-heading.pdf", raw)

    paras = [p for p in result.paragraphs if p.block_type == "paragraph"]
    assert len(paras) == 1
    assert "figyelembe veszi hogy" in paras[0].text
    assert "Tovabbi Feldolgozasi Lepesek" in paras[0].text


def test_extract_document_from_pdf_keeps_real_legal_heading_separate_from_body() -> None:
    raw = _build_pdf_real_legal_heading_followed_by_wrapped_paragraph()

    result = extract_document_from_upload("real-legal-heading.pdf", raw)

    assert len(result.paragraphs) == 1
    assert result.paragraphs[0].block_type == "heading"
    assert "9.C. A kockazatviseles idotartama" in result.paragraphs[0].text
    assert "Az idojaras-biztositas" in result.paragraphs[0].text
    assert "07:00 oraja." in result.paragraphs[0].text


def test_lonely_clause_number_line_is_not_classified_as_noise() -> None:
    assert _looks_like_lonely_clause_number_line("3.")
    assert _looks_like_lonely_clause_number_line("12.3.")
    assert not _looks_like_noise_line("3.")
    assert not _looks_like_noise_line("12.3.")


def test_repeated_long_edge_line_kept_when_only_two_occurrences() -> None:
    long_footer = (
        "This longer repeated disclaimer should not be removed when it appears only twice near the page edge."
    )
    ph = 800.0
    lines = [
        _PdfLine(
            text=long_footer,
            page_number=1,
            x0=72.0,
            top=10.0,
            x1=520.0,
            bottom=ph - 10.0,
            font_size=11.0,
            is_bold=False,
            page_height=ph,
            word_count=14,
            wide_gap_count=0,
            max_gap=0.0,
            cell_texts=(),
        ),
        _PdfLine(
            text=long_footer,
            page_number=2,
            x0=72.0,
            top=10.0,
            x1=520.0,
            bottom=ph - 10.0,
            font_size=11.0,
            is_bold=False,
            page_height=ph,
            word_count=14,
            wide_gap_count=0,
            max_gap=0.0,
            cell_texts=(),
        ),
    ]
    excluded = _detect_repeated_edge_lines(lines)
    assert excluded == set()


def test_pdf_lonely_number_line_merges_with_following_body() -> None:
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    _, height = A4

    pdf.setFont("Helvetica", 11)
    pdf.drawString(72, height - 100, "3.")
    pdf.drawString(90, height - 118, "A felek megallapodnak a feltetelekrol es a teljesitesrol.")
    pdf.save()
    raw = buffer.getvalue()

    result = extract_document_from_upload("lonely-three-dot.pdf", raw)

    list_items = [p for p in result.paragraphs if p.block_type == "list_item"]
    assert len(list_items) == 1
    assert list_items[0].text.startswith("3.")
    assert "feltetelekrol" in list_items[0].text


def test_extract_document_from_pdf_keeps_ellipsis_continuation_across_large_gap() -> None:
    raw = _build_pdf_ellipsis_then_title_case_next_line()

    result = extract_document_from_upload("ellipsis-flow.pdf", raw)

    paras = [p for p in result.paragraphs if p.block_type == "paragraph"]
    assert len(paras) == 1
    assert "..." in paras[0].text
    assert "Additional Title Case" in paras[0].text


def test_extract_document_from_pdf_merges_across_page_when_paragraph_flow_continues() -> None:
    raw = _build_pdf_cross_page_title_case_continuation()

    result = extract_document_from_upload("cross-page-flow.pdf", raw)

    paras = [paragraph for paragraph in result.paragraphs if paragraph.block_type == "paragraph"]
    assert len(paras) == 1
    assert paras[0].metadata.get("page_span") == [1, 2]
    assert "shall state that" in paras[0].text
    assert "Further Processing Steps" in paras[0].text


def test_extract_document_from_pdf_does_not_merge_across_page_after_sentence_end() -> None:
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    _, height = A4

    pdf.setFont("Helvetica", 11)
    pdf.drawString(72, height - 90, "First paragraph ends here completely.")
    pdf.showPage()
    pdf.setFont("Helvetica-Bold", 14)
    pdf.drawString(72, height - 90, "Real Section Title")
    pdf.save()
    raw = buffer.getvalue()

    result = extract_document_from_upload("page-break-section.pdf", raw)

    assert len([paragraph for paragraph in result.paragraphs if paragraph.block_type == "paragraph"]) == 1
    assert "Real Section Title" in result.paragraphs[0].text


def test_extract_document_from_docx_marks_heading_style_as_heading() -> None:
    raw = _build_docx_with_header()

    result = extract_document_from_upload("policy.docx", raw)

    assert result.metadata["source_format"] == "docx"
    assert result.paragraphs[0].block_type == "heading"
    assert result.paragraphs[0].metadata["style_name"] == "heading 1"
    assert result.paragraphs[1].block_type == "paragraph"


def test_extract_document_from_docx_marks_list_items() -> None:
    raw = _build_docx_with_list_items()

    result = extract_document_from_upload("list.docx", raw)

    list_items = [paragraph.text for paragraph in result.paragraphs if paragraph.block_type == "list_item"]
    assert list_items == [
        "- Elso listaelem a szabalyzatban",
        "2. Masodik listaelem kulon sorban",
    ]


def test_extract_document_from_pdf_marks_table_of_contents_as_metadata() -> None:
    raw = _build_table_of_contents_pdf()

    result = extract_document_from_upload("contents.pdf", raw)

    assert result.paragraphs
    assert result.paragraphs[0].block_type == "metadata"
    assert result.paragraphs[0].metadata["metadata_kind"] == "table_of_contents"
    assert "Tartalomjegyzék" in result.paragraphs[0].text


def test_extract_document_from_pdf_keeps_table_cells_in_metadata() -> None:
    raw = _build_sample_pdf()

    result = extract_document_from_upload("sample.pdf", raw)

    table_paragraph = next(paragraph for paragraph in result.paragraphs if paragraph.block_type == "table_row")
    assert table_paragraph.metadata["table_cells"] == ["Name", "Qty", "Price"]
    assert table_paragraph.metadata["table_role"] == "header"


def test_extract_document_from_docx_keeps_table_rows_and_headers() -> None:
    raw = _build_docx_with_table()

    result = extract_document_from_upload("table.docx", raw)

    assert any(
        paragraph.block_type == "metadata" and paragraph.metadata.get("metadata_kind") == "table_of_contents"
        for paragraph in result.paragraphs
    )
    table_rows = [paragraph for paragraph in result.paragraphs if paragraph.block_type == "table_row"]
    assert len(table_rows) == 2
    assert table_rows[0].metadata["table_role"] == "header"
    assert table_rows[1].metadata["table_column_headers"] == ["Mező", "Érték"]
    assert table_rows[1].metadata["table_cells"] == ["Bonus-malus", "A10"]
