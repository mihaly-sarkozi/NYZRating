# backend/shared/documents/text_extraction.py
# Feladat: Feltöltött dokumentumok általános TXT/PDF/DOCX szöveg- és struktúra-kinyerését vezérli. TXT esetén egyszerű dekódolást, PDF esetén layout parser delegálást, DOCX esetén bekezdés/tábla/heading/lista heuristikát használ, majd ExtractedDocument contractot ad vissza. Shared utility réteg, amelyet appok használhatnak dokumentum ingesthez.
# Sárközi Mihály - 2026.05.21

from __future__ import annotations

import io
import re
from dataclasses import replace
from statistics import median

from shared.documents.models import ExtractedDocument, ExtractedParagraph
from shared.documents.pdf_layout_parser import (
    _looks_like_lonely_clause_number_line,
    extract_pdf_layout,
)


def _looks_like_docx_header(text: str, *, is_bold: bool, font_size: float | None, baseline_font_size: float) -> bool:
    words = re.findall(r"\b\w+\b", text, flags=re.UNICODE)
    if not words:
        return False
    if len(words) > 18 or len(text) > 140:
        return False
    if re.match(r"^\s*(?:[-*•]|\d+(?:\.\d+){0,5}\.?)\s+\S+", text):
        return False
    if text.endswith((".", ";", ":")):
        return False
    return is_bold or bool(font_size and font_size >= baseline_font_size * 1.18)


def _looks_like_docx_list_item(text: str, *, style_name: str) -> bool:
    if style_name.startswith("list"):
        return True
    return bool(re.match(r"^\s*(?:[-*•]|\d+(?:\.\d+){0,5}\.?|[a-zA-Z]\))\s+\S+", text))


def _looks_like_table_of_contents(text: str) -> bool:
    normalized = re.sub(r"\s+", " ", (text or "").strip().lower())
    if normalized in {"tartalomjegyzék", "contents", "table of contents"}:
        return True
    return bool(
        re.match(
            r"^(?:\d+(?:\.\d+){0,5}\.?\s+)?[a-záéíóöőúüű0-9].*\.{2,}\s*\d+$",
            normalized,
        )
    )


def _looks_like_noise_line(text: str) -> bool:
    normalized = (text or "").strip()
    if not normalized:
        return True
    if re.fullmatch(r"[\W\d_]+", normalized, flags=re.UNICODE):
        return True
    if len(normalized) <= 40 and normalized.count(".") >= 6 and len(re.findall(r"\b\w+\b", normalized, flags=re.UNICODE)) <= 2:
        return True
    alpha_count = len(re.findall(r"[A-Za-zÁÉÍÓÖŐÚÜŰáéíóöőúüű]", normalized))
    return alpha_count == 0 and len(normalized) <= 60


def _looks_like_table_header_cells(cells: list[str]) -> bool:
    normalized_cells = [cell.strip() for cell in cells if cell.strip()]
    if len(normalized_cells) < 2:
        return False
    if any(len(cell) > 40 for cell in normalized_cells):
        return False
    digit_heavy_count = sum(any(char.isdigit() for char in cell) for cell in normalized_cells)
    return digit_heavy_count <= max(1, len(normalized_cells) // 3)


def _annotate_table_blocks(paragraphs: list[ExtractedParagraph]) -> list[ExtractedParagraph]:
    annotated: list[ExtractedParagraph] = []
    current_headers: list[str] | None = None
    previous_block_type: str | None = None
    for paragraph in paragraphs:
        if paragraph.block_type != "table_row":
            current_headers = None
            previous_block_type = paragraph.block_type
            annotated.append(paragraph)
            continue
        metadata = dict(paragraph.metadata or {})
        table_cells = [str(cell).strip() for cell in metadata.get("table_cells") or [] if str(cell).strip()]
        if previous_block_type != "table_row":
            current_headers = None
        if table_cells and current_headers is None and _looks_like_table_header_cells(table_cells):
            metadata["table_role"] = "header"
            metadata["table_column_headers"] = list(table_cells)
            current_headers = table_cells
        else:
            metadata["table_role"] = "row" if table_cells else "unknown"
            if current_headers:
                metadata["table_column_headers"] = list(current_headers)
        annotated.append(replace(paragraph, metadata=metadata))
        previous_block_type = paragraph.block_type
    return annotated


def _annotate_special_blocks(paragraphs: list[ExtractedParagraph]) -> list[ExtractedParagraph]:
    annotated: list[ExtractedParagraph] = []
    for paragraph in paragraphs:
        metadata = dict(paragraph.metadata or {})
        if paragraph.block_type == "metadata":
            metadata["metadata_kind"] = "table_of_contents"
        annotated.append(replace(paragraph, metadata=metadata))
    return annotated


def _iter_docx_blocks(document):
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    for child in document.element.body.iterchildren():
        if child.tag.endswith("}p"):
            yield "paragraph", Paragraph(child, document)
        elif child.tag.endswith("}tbl"):
            yield "table", Table(child, document)


def extract_document_from_upload(filename: str, raw: bytes) -> ExtractedDocument:
    name = (filename or "").lower()
    if name.endswith(".txt"):
        text = raw.decode("utf-8", errors="replace")
        return ExtractedDocument(
            text_content=text,
            paragraphs=[ExtractedParagraph(text=text)] if text.strip() else [],
            metadata={"source_format": "txt", "extraction_engine": "plain_text_v1"},
        )

    if name.endswith(".pdf"):
        return extract_pdf_layout(raw)

    if name.endswith(".docx"):
        from docx import Document

        document = Document(io.BytesIO(raw))
        font_sizes = [
            float(run.font.size.pt)
            for paragraph in document.paragraphs
            for run in paragraph.runs
            if run.font.size is not None
        ]
        baseline_font_size = median(font_sizes) if font_sizes else 11.0
        paragraphs: list[ExtractedParagraph] = []
        for block_kind, block in _iter_docx_blocks(document):
            if block_kind == "paragraph":
                text = (block.text or "").strip()
                if not text:
                    continue
                run_sizes = [float(run.font.size.pt) for run in block.runs if run.font.size is not None]
                font_size = median(run_sizes) if run_sizes else None
                is_bold = any(bool(run.bold) for run in block.runs if (run.text or "").strip())
                style_name = str(getattr(block.style, "name", "") or "").strip().lower()
                block_type = "paragraph"
                if style_name.startswith("toc") or _looks_like_table_of_contents(text):
                    block_type = "metadata"
                elif _looks_like_lonely_clause_number_line(text):
                    block_type = "list_item"
                elif _looks_like_docx_list_item(text, style_name=style_name):
                    block_type = "list_item"
                elif style_name.startswith("heading") or style_name in {"title", "subtitle"}:
                    block_type = "heading"
                elif _looks_like_docx_header(text, is_bold=is_bold, font_size=font_size, baseline_font_size=baseline_font_size):
                    block_type = "heading"
                paragraphs.append(
                    ExtractedParagraph(
                        text=text,
                        block_type=block_type,
                        font_size=round(font_size, 2) if font_size is not None else None,
                        is_bold=is_bold,
                        metadata={"style_name": style_name or None, "block_type": block_type},
                    )
                )
                continue

            for row_index, row in enumerate(block.rows):
                cells = [re.sub(r"\s+", " ", (cell.text or "").strip()) for cell in row.cells]
                cells = [cell for cell in cells if cell]
                if not cells:
                    continue
                paragraphs.append(
                    ExtractedParagraph(
                        text=" | ".join(cells),
                        block_type="table_row",
                        metadata={
                            "block_type": "table_row",
                            "table_cells": cells,
                            "docx_table_row_index": row_index,
                            "docx_table_column_count": len(cells),
                        },
                    )
                )
        paragraphs = _annotate_special_blocks(_annotate_table_blocks(paragraphs))
        return ExtractedDocument(
            text_content="\n\n".join(paragraph.text for paragraph in paragraphs),
            paragraphs=paragraphs,
            metadata={
                "source_format": "docx",
                "extraction_engine": "python_docx_v1",
                "baseline_font_size": round(float(baseline_font_size), 2),
            },
        )

    raise ValueError("unsupported_file_type")


def extract_text_from_upload(filename: str, raw: bytes) -> str:
    return extract_document_from_upload(filename, raw).text_content
